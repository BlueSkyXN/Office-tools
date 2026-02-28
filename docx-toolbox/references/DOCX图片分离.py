#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX 图片分离工具 - 增强版 v2.0
将 DOCX 文件中的图片提取出来并生成独立的 PDF 文件，同时在原文档中标记图片位置

增强功能:
1. 智能连续编号：只对实际使用的图片进行连续编号
2. PDF多页目录：在PDF开头生成图片索引目录（A3纵向，支持多页）
3. 孤儿图片检测：识别并在PDF末尾列出未被引用的图片
4. 全面位置追踪：记录每张图片在文档中的确切位置
   - 正文段落和表格
   - 页眉和页脚（所有节）
   - 文本框和形状
5. 格式兼容性：支持 DrawingML 和 VML（旧版）格式的图片

v2.0 更新:
- 修复页眉、页脚、文本框中的图片被误判为"孤儿图片"的问题
- 增加对旧版 VML 格式图片的支持
- 改进位置描述的准确性和可读性

用法:
  python DOCX图片分离.py <docx文件路径|文件夹路径> [--remove-images] [--output-dir <输出目录>]
"""

import sys
import os
from pathlib import Path
import argparse
from docx import Document
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsmap
from PIL import Image

# 注册 VML 命名空间（如果尚未注册）
if 'v' not in nsmap:
    nsmap['v'] = 'urn:schemas-microsoft-com:vml'
if 'o' not in nsmap:
    nsmap['o'] = 'urn:schemas-microsoft-com:office:office'
import io
from reportlab.lib.pagesizes import A4, A3
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.colors import HexColor


# A4 尺寸 (单位: 点)
A4_WIDTH, A4_HEIGHT = A4  # 595.28 x 841.89 点

# A3 纵向尺寸 (单位: 点)
A3_WIDTH, A3_HEIGHT = A3  # 841.89 x 1190.55 点


def register_fonts():
    """注册中文字体"""
    try:
        # macOS 系统字体
        if os.path.exists('/System/Library/Fonts/STHeiti Light.ttc'):
            pdfmetrics.registerFont(TTFont('Chinese', '/System/Library/Fonts/STHeiti Light.ttc'))
            return 'Chinese'
        elif os.path.exists('/System/Library/Fonts/PingFang.ttc'):
            pdfmetrics.registerFont(TTFont('Chinese', '/System/Library/Fonts/PingFang.ttc'))
            return 'Chinese'
        # Windows 系统字体
        elif os.path.exists('C:\\Windows\\Fonts\\simhei.ttf'):
            pdfmetrics.registerFont(TTFont('Chinese', 'C:\\Windows\\Fonts\\simhei.ttf'))
            return 'Chinese'
    except Exception as e:
        print(f"  ⚠️  中文字体注册失败，使用默认字体: {e}")
    return 'Helvetica'


def _extract_images_from_element(element, location_prefix, all_images, used_rids, active_images):
    """
    从 XML 元素中提取图片引用（支持 DrawingML 和 VML）

    Args:
        element: XML 元素
        location_prefix: 位置前缀描述
        all_images: 所有图片资源字典
        used_rids: 已使用的 rId 集合
        active_images: 活跃图片列表
    """
    # 1. 查找 DrawingML 格式 (w:drawing)
    for drawing in element.findall('.//' + qn('w:drawing')):
        blip = drawing.find('.//' + qn('a:blip'))
        if blip is not None:
            embed_id = blip.get(qn('r:embed'))
            if embed_id and embed_id in all_images:
                if embed_id not in used_rids:  # 避免重复
                    used_rids.add(embed_id)
                    active_images.append((embed_id, location_prefix))

    # 2. 查找 VML 格式 (w:pict -> v:shape -> v:imagedata)
    for pict in element.findall('.//' + qn('w:pict')):
        # VML 图片数据在 v:imagedata 标签中
        for imagedata in pict.findall('.//' + qn('v:imagedata')):
            # VML 使用 r:id 属性引用图片
            embed_id = imagedata.get(qn('r:id'))
            if not embed_id:
                # 有些 VML 使用 o:relid
                embed_id = imagedata.get(qn('o:relid'))
            if embed_id and embed_id in all_images:
                if embed_id not in used_rids:
                    used_rids.add(embed_id)
                    active_images.append((embed_id, f"{location_prefix}[VML]"))


def analyze_document_images(docx_path):
    """
    分析文档中的图片使用情况（增强版）

    扫描范围:
    - 正文段落和表格
    - 页眉和页脚（所有节）
    - 文本框和形状
    - 支持 DrawingML 和 VML 格式

    Returns:
        dict: {
            'active_images': [(rel_id, location_info), ...],  # 被引用的图片
            'orphan_images': [rel_id, ...],  # 孤儿图片
            'all_images': {rel_id: image_info, ...}  # 所有图片资源
        }
    """
    doc = Document(docx_path)

    # 收集所有图片资源（包括主文档和页眉页脚的 relationship）
    all_images = {}

    # 从主文档收集图片
    for rel in doc.part.rels.values():
        # 检查 reltype 而不是 target_ref，因为图片的 target_ref 可能是 "media/rId50.png"
        if "image" in rel.reltype:
            try:
                content_type = rel.target_part.content_type
                image_format = content_type.split('/')[-1]
                if image_format == 'jpeg':
                    image_format = 'jpg'

                all_images[rel.rId] = {
                    'data': rel.target_part.blob,
                    'format': image_format,
                    'size': len(rel.target_part.blob),
                    'target': rel.target_ref
                }
            except Exception as e:
                print(f"  ⚠️  跳过无效图片关系 {rel.rId}: {e}")

    # 从页眉/页脚收集图片
    for section in doc.sections:
        # 页眉
        try:
            header = section.header
            if header and hasattr(header, 'part') and hasattr(header.part, 'rels'):
                for rel in header.part.rels.values():
                    if "image" in rel.reltype and rel.rId not in all_images:
                        try:
                            content_type = rel.target_part.content_type
                            image_format = content_type.split('/')[-1]
                            if image_format == 'jpeg':
                                image_format = 'jpg'

                            all_images[rel.rId] = {
                                'data': rel.target_part.blob,
                                'format': image_format,
                                'size': len(rel.target_part.blob),
                                'target': rel.target_ref
                            }
                        except Exception:
                            pass
        except Exception:
            pass

        # 页脚
        try:
            footer = section.footer
            if footer and hasattr(footer, 'part') and hasattr(footer.part, 'rels'):
                for rel in footer.part.rels.values():
                    if "image" in rel.reltype and rel.rId not in all_images:
                        try:
                            content_type = rel.target_part.content_type
                            image_format = content_type.split('/')[-1]
                            if image_format == 'jpeg':
                                image_format = 'jpg'

                            all_images[rel.rId] = {
                                'data': rel.target_part.blob,
                                'format': image_format,
                                'size': len(rel.target_part.blob),
                                'target': rel.target_ref
                            }
                        except Exception:
                            pass
        except Exception:
            pass

    # 收集被引用的图片及其位置
    active_images = []
    used_rids = set()

    # 1. 扫描正文段落
    for para_idx, paragraph in enumerate(doc.paragraphs):
        # 获取段落文本预览
        text_preview = paragraph.text.strip()[:50]
        if text_preview:
            location = f"正文-段落{para_idx}: {text_preview}"
        else:
            location = f"正文-段落{para_idx}"

        _extract_images_from_element(
            paragraph._element,
            location,
            all_images,
            used_rids,
            active_images
        )

    # 2. 扫描正文表格
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    location = f"正文-表格{table_idx+1}-行{row_idx+1}-列{cell_idx+1}"
                    _extract_images_from_element(
                        para._element,
                        location,
                        all_images,
                        used_rids,
                        active_images
                    )

    # 3. 扫描页眉和页脚（所有节）
    for section_idx, section in enumerate(doc.sections):
        # 3.1 页眉（首页、偶数页、默认）
        for header_type, header_name in [
            (1, "首页页眉"),  # WD_HEADER_FOOTER.PRIMARY
            (2, "偶数页页眉"),  # WD_HEADER_FOOTER.EVEN_PAGE
            (3, "默认页眉")   # WD_HEADER_FOOTER.FIRST_PAGE
        ]:
            try:
                header = section.header
                if header:
                    # 扫描页眉段落
                    for para in header.paragraphs:
                        location = f"第{section_idx+1}节-{header_name}"
                        _extract_images_from_element(
                            para._element,
                            location,
                            all_images,
                            used_rids,
                            active_images
                        )

                    # 扫描页眉表格
                    for table_idx, table in enumerate(header.tables):
                        for row_idx, row in enumerate(table.rows):
                            for cell_idx, cell in enumerate(row.cells):
                                for para in cell.paragraphs:
                                    location = f"第{section_idx+1}节-{header_name}-表格{table_idx+1}"
                                    _extract_images_from_element(
                                        para._element,
                                        location,
                                        all_images,
                                        used_rids,
                                        active_images
                                    )
            except Exception as e:
                # 某些文档可能没有特定类型的页眉
                pass

        # 3.2 页脚（首页、偶数页、默认）
        for footer_type, footer_name in [
            (1, "首页页脚"),
            (2, "偶数页页脚"),
            (3, "默认页脚")
        ]:
            try:
                footer = section.footer
                if footer:
                    # 扫描页脚段落
                    for para in footer.paragraphs:
                        location = f"第{section_idx+1}节-{footer_name}"
                        _extract_images_from_element(
                            para._element,
                            location,
                            all_images,
                            used_rids,
                            active_images
                        )

                    # 扫描页脚表格
                    for table_idx, table in enumerate(footer.tables):
                        for row_idx, row in enumerate(table.rows):
                            for cell_idx, cell in enumerate(row.cells):
                                for para in cell.paragraphs:
                                    location = f"第{section_idx+1}节-{footer_name}-表格{table_idx+1}"
                                    _extract_images_from_element(
                                        para._element,
                                        location,
                                        all_images,
                                        used_rids,
                                        active_images
                                    )
            except Exception as e:
                pass

    # 4. 递归扫描文本框和形状（通过 XML 底层遍历）
    # 文本框内容存储在 w:txbxContent 标签中
    try:
        body_element = doc.element.body
        for txbx_content in body_element.findall('.//' + qn('w:txbxContent')):
            # 在文本框内查找段落
            for para_idx, para_element in enumerate(txbx_content.findall('.//' + qn('w:p'))):
                location = f"文本框-段落{para_idx}"
                _extract_images_from_element(
                    para_element,
                    location,
                    all_images,
                    used_rids,
                    active_images
                )
    except Exception as e:
        # 某些文档可能没有文本框
        pass

    # 识别孤儿图片
    orphan_images = [rid for rid in all_images.keys() if rid not in used_rids]

    return {
        'active_images': active_images,
        'orphan_images': orphan_images,
        'all_images': all_images
    }


def create_catalog_pages(c, analysis_result, font_name):
    """
    在PDF中创建多页目录（A3纵向）

    Args:
        c: canvas对象
        analysis_result: 文档分析结果
        font_name: 字体名称
    """
    active_images = analysis_result['active_images']
    active_count = len(active_images)
    orphan_count = len(analysis_result['orphan_images'])
    total_count = len(analysis_result['all_images'])

    # 第一页：标题和统计
    c.setPageSize((A3_WIDTH, A3_HEIGHT))

    # 标题
    c.setFont(font_name, 24)
    c.setFillColor(HexColor('#1a1a1a'))
    c.drawString(60, A3_HEIGHT - 60, "图片索引目录")

    # 分隔线
    c.setStrokeColor(HexColor('#cccccc'))
    c.setLineWidth(1)
    c.line(60, A3_HEIGHT - 75, A3_WIDTH - 60, A3_HEIGHT - 75)

    # 统计信息
    c.setFont(font_name, 14)
    c.setFillColor(HexColor('#333333'))

    y = A3_HEIGHT - 110
    c.drawString(60, y, f"文档统计:")
    y -= 30
    c.setFont(font_name, 12)
    c.drawString(80, y, f"• 总图片数: {total_count} 张")
    y -= 25
    c.drawString(80, y, f"• 有效图片: {active_count} 张 (已标记并连续编号)")
    y -= 25
    c.drawString(80, y, f"• 孤儿图片: {orphan_count} 张 (未被引用)")

    # 有效图片索引标题
    y -= 50
    c.setFont(font_name, 14)
    c.setFillColor(HexColor('#1a1a1a'))
    c.drawString(60, y, "有效图片索引:")

    y -= 35
    c.setFont(font_name, 10)
    c.setFillColor(HexColor('#555555'))

    # 分页显示所有图片索引
    page_num = 1
    items_per_page_first = 35  # 第一页显示35个（留空间给标题）
    items_per_page_rest = 50   # 后续页每页显示50个

    for idx, (rel_id, location) in enumerate(active_images, 1):
        # 检查是否需要换页
        if page_num == 1 and idx > items_per_page_first:
            # 页脚
            c.setFont(font_name, 8)
            c.setFillColor(HexColor('#999999'))
            c.drawString(60, 40, f"生成工具: DOCX图片分离工具 - 增强版")
            c.drawString(A3_WIDTH - 250, 40, f"目录第 {page_num} 页")
            c.showPage()

            # 新页面
            c.setPageSize((A3_WIDTH, A3_HEIGHT))
            page_num += 1
            y = A3_HEIGHT - 60
            c.setFont(font_name, 14)
            c.setFillColor(HexColor('#1a1a1a'))
            c.drawString(60, y, f"有效图片索引 (续):")
            y -= 35
            c.setFont(font_name, 10)
            c.setFillColor(HexColor('#555555'))
        elif page_num > 1 and (idx - items_per_page_first - 1) % items_per_page_rest == 0 and idx > items_per_page_first:
            # 页脚
            c.setFont(font_name, 8)
            c.setFillColor(HexColor('#999999'))
            c.drawString(60, 40, f"生成工具: DOCX图片分离工具 - 增强版")
            c.drawString(A3_WIDTH - 250, 40, f"目录第 {page_num} 页")
            c.showPage()

            # 新页面
            c.setPageSize((A3_WIDTH, A3_HEIGHT))
            page_num += 1
            y = A3_HEIGHT - 60
            c.setFont(font_name, 14)
            c.setFillColor(HexColor('#1a1a1a'))
            c.drawString(60, y, f"有效图片索引 (续):")
            y -= 35
            c.setFont(font_name, 10)
            c.setFillColor(HexColor('#555555'))

        # 绘制图片索引
        if y < 80:  # 安全边距，不应该到达这里
            c.setFont(font_name, 8)
            c.setFillColor(HexColor('#999999'))
            c.drawString(60, 40, f"生成工具: DOCX图片分离工具 - 增强版")
            c.drawString(A3_WIDTH - 250, 40, f"目录第 {page_num} 页")
            c.showPage()

            c.setPageSize((A3_WIDTH, A3_HEIGHT))
            page_num += 1
            y = A3_HEIGHT - 60
            c.setFont(font_name, 14)
            c.setFillColor(HexColor('#1a1a1a'))
            c.drawString(60, y, f"有效图片索引 (续):")
            y -= 35
            c.setFont(font_name, 10)
            c.setFillColor(HexColor('#555555'))

        pdf_page = idx + page_num  # 目录页数 + 图片编号
        text = f"图{idx} → PDF第{pdf_page}页 | 位置: {location}"

        # 文本过长则截断
        if len(text) > 110:
            text = text[:107] + "..."

        c.drawString(70, y, text)
        y -= 20

    # 孤儿图片说明（在最后一页）
    if orphan_count > 0:
        y -= 30
        if y < 200:  # 空间不够，新开一页
            c.setFont(font_name, 8)
            c.setFillColor(HexColor('#999999'))
            c.drawString(60, 40, f"生成工具: DOCX图片分离工具 - 增强版")
            c.drawString(A3_WIDTH - 250, 40, f"目录第 {page_num} 页")
            c.showPage()

            c.setPageSize((A3_WIDTH, A3_HEIGHT))
            page_num += 1
            y = A3_HEIGHT - 60

        c.setFont(font_name, 14)
        c.setFillColor(HexColor('#d32f2f'))
        c.drawString(60, y, f"孤儿图片 ({orphan_count} 张):")

        y -= 30
        c.setFont(font_name, 11)
        c.setFillColor(HexColor('#666666'))
        orphan_start_page = active_count + page_num + 1
        c.drawString(70, y, f"未被文档引用的图片已附在PDF末尾 (第{orphan_start_page}页起)")
        y -= 22
        c.drawString(70, y, "这些图片可能是:")
        y -= 20
        c.drawString(85, y, "• 删除后残留的图片资源")
        y -= 20
        c.drawString(85, y, "• 被其他元素（页眉/页脚/文本框）引用")
        y -= 20
        c.drawString(85, y, "• 重复导入但未使用的图片")

    # 最后一页的页脚
    c.setFont(font_name, 8)
    c.setFillColor(HexColor('#999999'))
    c.drawString(60, 40, f"生成工具: DOCX图片分离工具 - 增强版")
    c.drawString(A3_WIDTH - 250, 40, f"目录第 {page_num} 页 / 共 {page_num} 页")

    c.showPage()
    return page_num  # 返回目录页数


def extract_images_from_docx(docx_path):
    """
    从 DOCX 文件中提取所有图片

    Args:
        docx_path: DOCX 文件路径

    Returns:
        list: [(image_data, image_format, image_index), ...]
    """
    doc = Document(docx_path)
    images = []
    image_index = 1

    # 遍历所有关系，找到图片
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            # 获取图片格式
            content_type = rel.target_part.content_type
            image_format = content_type.split('/')[-1]
            if image_format == 'jpeg':
                image_format = 'jpg'

            images.append({
                'data': image_data,
                'format': image_format,
                'index': image_index,
                'rel_id': rel.rId
            })
            image_index += 1

    return images


def get_image_runs_in_paragraph(paragraph):
    """
    获取段落中包含图片的 run

    Args:
        paragraph: python-docx 段落对象

    Returns:
        list: [(run, drawing_element), ...]
    """
    image_runs = []
    for run in paragraph.runs:
        # 查找 run 中的图片元素
        for drawing in run._element.findall('.//' + qn('w:drawing')):
            image_runs.append((run, drawing))
    return image_runs


def _mark_images_in_element(element, rel_id_to_index, remove_images):
    """
    在 XML 元素中标记图片（支持 DrawingML 和 VML）

    Args:
        element: XML 段落元素 (w:p)
        rel_id_to_index: rId 到编号的映射
        remove_images: 是否删除原图

    Returns:
        int: 标记的图片数量
    """
    count = 0

    # 1. 处理 DrawingML 格式 (w:drawing)
    # 需要遍历所有 run (w:r)
    for run_element in element.findall('./' + qn('w:r')):
        drawings = run_element.findall('.//' + qn('w:drawing'))
        for drawing in drawings:
            blip = drawing.find('.//' + qn('a:blip'))
            if blip is not None:
                embed_id = blip.get(qn('r:embed'))
                if embed_id in rel_id_to_index:
                    img_num = rel_id_to_index[embed_id]

                    # 插入标记
                    new_run_element = parse_xml(
                        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        f'<w:t>【图{img_num}】</w:t>'
                        f'</w:r>'
                    )

                    run_index = list(element).index(run_element)
                    element.insert(run_index, new_run_element)

                    if remove_images:
                        run_element.remove(drawing)

                    count += 1

    # 2. 处理 VML 格式 (w:pict)
    for run_element in element.findall('./' + qn('w:r')):
        picts = run_element.findall('.//' + qn('w:pict'))
        for pict in picts:
            for imagedata in pict.findall('.//' + qn('v:imagedata')):
                embed_id = imagedata.get(qn('r:id'))
                if embed_id in rel_id_to_index:
                    img_num = rel_id_to_index[embed_id]

                    # 插入标记
                    new_run_element = parse_xml(
                        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        f'<w:t>【图{img_num}】</w:t>'
                        f'</w:r>'
                    )

                    run_index = list(element).index(run_element)
                    element.insert(run_index, new_run_element)

                    if remove_images:
                        run_element.remove(pict)

                    count += 1

    return count


def mark_images_in_docx(docx_path, output_path, analysis_result, remove_images=False):
    """
    在DOCX中标记图片（增强版：支持页眉、页脚、文本框和 VML）

    Args:
        docx_path: 输入文件路径
        output_path: 输出文件路径
        analysis_result: 文档分析结果
        remove_images: 是否删除原图
    """
    doc = Document(docx_path)

    # 建立 rId 到连续编号的映射
    rel_id_to_index = {}
    for idx, (rel_id, location) in enumerate(analysis_result['active_images'], 1):
        # 避免重复的 rel_id
        if rel_id not in rel_id_to_index:
            rel_id_to_index[rel_id] = idx

    replaced_count = 0

    # 1. 处理正文段落
    for paragraph in doc.paragraphs:
        count = _mark_images_in_element(
            paragraph._element,
            rel_id_to_index,
            remove_images
        )
        replaced_count += count

    # 2. 处理正文表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    count = _mark_images_in_element(
                        paragraph._element,
                        rel_id_to_index,
                        remove_images
                    )
                    replaced_count += count

    # 3. 处理页眉和页脚
    for section in doc.sections:
        # 3.1 页眉
        try:
            header = section.header
            if header:
                # 页眉段落
                for paragraph in header.paragraphs:
                    count = _mark_images_in_element(
                        paragraph._element,
                        rel_id_to_index,
                        remove_images
                    )
                    replaced_count += count

                # 页眉表格
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                count = _mark_images_in_element(
                                    paragraph._element,
                                    rel_id_to_index,
                                    remove_images
                                )
                                replaced_count += count
        except Exception as e:
            pass

        # 3.2 页脚
        try:
            footer = section.footer
            if footer:
                # 页脚段落
                for paragraph in footer.paragraphs:
                    count = _mark_images_in_element(
                        paragraph._element,
                        rel_id_to_index,
                        remove_images
                    )
                    replaced_count += count

                # 页脚表格
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                count = _mark_images_in_element(
                                    paragraph._element,
                                    rel_id_to_index,
                                    remove_images
                                )
                                replaced_count += count
        except Exception as e:
            pass

    # 4. 处理文本框
    try:
        body_element = doc.element.body
        for txbx_content in body_element.findall('.//' + qn('w:txbxContent')):
            for para_element in txbx_content.findall('.//' + qn('w:p')):
                count = _mark_images_in_element(
                    para_element,
                    rel_id_to_index,
                    remove_images
                )
                replaced_count += count
    except Exception as e:
        pass

    doc.save(output_path)
    return replaced_count


def calculate_page_size(image_width, image_height):
    """
    计算 PDF 页面大小

    规则:
    - 小于 A4 的图片等比例放大，直到至少一边达到 A4，页面大小为放大后的图片大小
    - 大于等于 A4 的图片使用原图大小

    Args:
        image_width: 图片宽度（像素）
        image_height: 图片高度（像素）

    Returns:
        (page_width, page_height, scale): 页面宽度、高度（点）和缩放比例
    """
    # 假设图片 DPI 为 72（PDF 默认）
    img_width_pt = image_width
    img_height_pt = image_height

    # 如果图片两边都大于等于 A4，使用原图尺寸
    if img_width_pt >= A4_WIDTH and img_height_pt >= A4_HEIGHT:
        return img_width_pt, img_height_pt, 1.0

    # 如果图片至少有一边小于 A4，需要放大
    # 计算宽度和高度需要的缩放比例
    width_scale = A4_WIDTH / img_width_pt
    height_scale = A4_HEIGHT / img_height_pt

    # 选择较小的缩放比例，确保等比例放大且至少一边达到 A4
    scale = min(width_scale, height_scale)

    # 计算放大后的页面尺寸（等比例）
    page_width = img_width_pt * scale
    page_height = img_height_pt * scale

    return page_width, page_height, scale


def optimize_image_for_pdf(img, original_format='png', quality=85):
    """
    优化图片以减小PDF大小

    Args:
        img: PIL Image对象
        original_format: 原始格式 ('png', 'jpeg')
        quality: JPEG质量 (1-100)

    Returns:
        (img_buffer, format): 优化后的图片数据和格式
    """
    img_buffer = io.BytesIO()

    # 转换RGBA为RGB（JPEG不支持透明）
    if img.mode == 'RGBA':
        # 检查是否真的有透明通道
        alpha = img.split()[3]
        if alpha.getextrema() == (255, 255):
            # 没有透明，可以安全转JPEG
            rgb_img = Image.new('RGB', img.size, (255, 255, 255))
            rgb_img.paste(img, mask=img.split()[3])
            img = rgb_img
            use_jpeg = True
        else:
            # 有透明，必须用PNG
            use_jpeg = False
    elif img.mode == 'RGB':
        use_jpeg = True
    elif img.mode == 'L':
        # 灰度图
        use_jpeg = True
    else:
        # 其他模式转RGB
        img = img.convert('RGB')
        use_jpeg = True

    # 如果原始是JPEG且无透明，保持JPEG
    if original_format == 'jpeg' and use_jpeg:
        img.save(img_buffer, format='JPEG', quality=quality, optimize=True)
        return img_buffer, 'JPEG'

    # 对于PNG，检查是否应该转JPEG
    if use_jpeg and img.mode in ['RGB', 'L']:
        # 计算复杂度（简单方法：检查颜色数量）
        # 如果是截图/图表（颜色少），用PNG；如果是照片（颜色多），用JPEG
        try:
            colors_result = img.getcolors(maxcolors=256)
            if colors_result is None:
                # 颜色超过256种，可能是照片，用JPEG
                img.save(img_buffer, format='JPEG', quality=quality, optimize=True)
                return img_buffer, 'JPEG'
            elif len(colors_result) > 128:
                # 颜色多，可能是照片，用JPEG
                img.save(img_buffer, format='JPEG', quality=quality, optimize=True)
                return img_buffer, 'JPEG'
            else:
                # 颜色少，可能是图表/截图，用PNG
                img.save(img_buffer, format='PNG', optimize=True)
                return img_buffer, 'PNG'
        except Exception:
            # 出错时默认用JPEG
            img.save(img_buffer, format='JPEG', quality=quality, optimize=True)
            return img_buffer, 'JPEG'
    else:
        # 保持PNG（有透明或其他原因）
        img.save(img_buffer, format='PNG', optimize=True)
        return img_buffer, 'PNG'


def create_pdf_with_catalog(analysis_result, output_pdf_path, optimize=True, jpeg_quality=85):
    """
    创建带多页目录的PDF文件

    Args:
        analysis_result: 文档分析结果
        output_pdf_path: 输出PDF路径
        optimize: 是否优化图片格式
        jpeg_quality: JPEG质量 (1-100)
    """
    all_images = analysis_result['all_images']
    active_images = analysis_result['active_images']
    orphan_images = analysis_result['orphan_images']

    if not all_images:
        print("⚠️  没有图片可以生成 PDF")
        return False

    # 注册字体
    font_name = register_fonts()

    # 创建PDF
    c = canvas.Canvas(str(output_pdf_path), pagesize=(A3_WIDTH, A3_HEIGHT))

    # 生成多页目录
    print("  📑 生成目录页...")
    catalog_pages = create_catalog_pages(c, analysis_result, font_name)
    print(f"     目录共 {catalog_pages} 页")

    # 添加有效图片（连续编号）
    print(f"\n  📸 添加有效图片 ({len(active_images)} 张):")
    if optimize:
        print(f"     优化模式: JPEG质量={jpeg_quality}, 智能格式选择")

    total_original_size = 0
    total_optimized_size = 0

    for idx, (rel_id, location) in enumerate(active_images, 1):
        try:
            img_info = all_images[rel_id]
            image_data = img_info['data']
            original_size = len(image_data)
            total_original_size += original_size

            img = Image.open(io.BytesIO(image_data))
            img_width, img_height = img.size
            page_width, page_height, scale = calculate_page_size(img_width, img_height)

            c.setPageSize((page_width, page_height))

            # 优化图片
            if optimize:
                img_buffer, final_format = optimize_image_for_pdf(
                    img,
                    original_format=img_info['format'],
                    quality=jpeg_quality
                )
            else:
                # 不优化，转PNG
                if img.mode == 'RGBA':
                    rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                    rgb_img.paste(img, mask=img.split()[3])
                    img = rgb_img
                elif img.mode not in ['RGB', 'L']:
                    img = img.convert('RGB')

                img_buffer = io.BytesIO()
                img.save(img_buffer, format='PNG')
                final_format = 'PNG'

            img_buffer.seek(0)
            optimized_size = len(img_buffer.getvalue())
            total_optimized_size += optimized_size

            img_reader = ImageReader(img_buffer)

            x = (page_width - img_width * scale) / 2
            y = (page_height - img_height * scale) / 2

            c.drawImage(img_reader, x, y, width=img_width * scale, height=img_height * scale)

            # 添加页面标注（右上角）
            c.setFont(font_name, 8)
            c.setFillColor(HexColor('#666666'))
            page_label = f"图{idx} | {location[:40]}"
            c.drawString(10, page_height - 15, page_label)

            c.showPage()

            # 显示优化信息
            if optimize:
                ratio = (1 - optimized_size/original_size) * 100
                print(f"    ✓ 图{idx}: {img_width}x{img_height}px | {final_format} | "
                      f"{original_size//1024}KB→{optimized_size//1024}KB ({ratio:+.0f}%) | {location[:30]}")
            else:
                print(f"    ✓ 图{idx}: {img_width}x{img_height}px | {location[:50]}")

        except Exception as e:
            print(f"    ❌ 图{idx} 处理失败: {e}")
            continue

    # 添加孤儿图片（如果有）
    if orphan_images:
        print(f"\n  🗑️  添加孤儿图片 ({len(orphan_images)} 张):")
        for idx, rel_id in enumerate(orphan_images, 1):
            try:
                img_info = all_images[rel_id]
                image_data = img_info['data']
                original_size = len(image_data)
                total_original_size += original_size

                img = Image.open(io.BytesIO(image_data))
                img_width, img_height = img.size
                page_width, page_height, scale = calculate_page_size(img_width, img_height)

                c.setPageSize((page_width, page_height))

                # 优化图片
                if optimize:
                    img_buffer, final_format = optimize_image_for_pdf(
                        img,
                        original_format=img_info['format'],
                        quality=jpeg_quality
                    )
                else:
                    if img.mode == 'RGBA':
                        rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                        rgb_img.paste(img, mask=img.split()[3])
                        img = rgb_img
                    elif img.mode not in ['RGB', 'L']:
                        img = img.convert('RGB')

                    img_buffer = io.BytesIO()
                    img.save(img_buffer, format='PNG')
                    final_format = 'PNG'

                img_buffer.seek(0)
                optimized_size = len(img_buffer.getvalue())
                total_optimized_size += optimized_size

                img_reader = ImageReader(img_buffer)

                x = (page_width - img_width * scale) / 2
                y = (page_height - img_height * scale) / 2

                c.drawImage(img_reader, x, y, width=img_width * scale, height=img_height * scale)

                # 添加"未使用"标注
                c.setFont(font_name, 10)
                c.setFillColor(HexColor('#d32f2f'))
                c.drawString(10, page_height - 15, f"[未使用] {rel_id} | {img_info['size']//1024}KB")

                c.showPage()

                if optimize:
                    ratio = (1 - optimized_size/original_size) * 100
                    print(f"    • {rel_id}: {img_width}x{img_height}px | {final_format} | "
                          f"{original_size//1024}KB→{optimized_size//1024}KB ({ratio:+.0f}%)")
                else:
                    print(f"    • {rel_id}: {img_width}x{img_height}px ({img_info['size']//1024}KB)")

            except Exception as e:
                print(f"    ❌ {rel_id} 处理失败: {e}")
                continue

    c.save()

    # 显示优化统计
    if optimize and total_original_size > 0:
        compression_ratio = (1 - total_optimized_size/total_original_size) * 100
        print(f"\n📊 图片优化统计:")
        print(f"   原始总大小: {total_original_size/1024/1024:.2f} MB")
        print(f"   优化后大小: {total_optimized_size/1024/1024:.2f} MB")
        print(f"   压缩率: {compression_ratio:.1f}%")

    print(f"\n✅ PDF 生成成功: {output_pdf_path}")
    return True


def process_docx_file(docx_path, remove_images=False, output_dir=None, optimize_images=True, jpeg_quality=85):
    """
    处理单个 DOCX 文件

    Args:
        docx_path: DOCX 文件路径
        remove_images: 是否删除原文档中的图片
        output_dir: 输出目录
        optimize_images: 是否优化图片格式
        jpeg_quality: JPEG质量 (1-100)
    """
    docx_path = Path(docx_path)

    if not docx_path.exists():
        print(f"❌ 文件不存在: {docx_path}")
        return False

    if not docx_path.suffix.lower() == '.docx':
        print(f"❌ 不是 DOCX 文件: {docx_path}")
        return False

    print(f"\n📄 处理文件: {docx_path.name}")
    print("=" * 80)

    # 确定输出目录
    if output_dir:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
    else:
        output_dir = docx_path.parent

    base_name = docx_path.stem
    output_docx_path = output_dir / f"{base_name}_已标记图片.docx"
    output_pdf_path = output_dir / f"{base_name}_附图.pdf"

    try:
        # 1. 分析文档图片
        print("🔍 分析文档图片...")
        analysis_result = analyze_document_images(docx_path)

        active_count = len(analysis_result['active_images'])
        orphan_count = len(analysis_result['orphan_images'])
        total_count = len(analysis_result['all_images'])

        print(f"  • 图片资源总数: {total_count}")
        print(f"  • 有效图片: {active_count} (被文档引用)")
        print(f"  • 孤儿图片: {orphan_count} (未被引用)")

        if total_count == 0:
            print("⚠️  未找到图片")
            return False

        # 2. 生成带目录的PDF
        print("\n📚 生成PDF...")
        create_pdf_with_catalog(analysis_result, output_pdf_path, optimize=optimize_images, jpeg_quality=jpeg_quality)

        # 3. 标记DOCX
        print("\n🏷️  标记图片位置...")
        replaced_count = mark_images_in_docx(
            docx_path,
            output_docx_path,
            analysis_result,
            remove_images
        )
        print(f"  ✓ 已标记 {replaced_count} 张图片 (连续编号: 图1~图{active_count})")

        if remove_images:
            print(f"  ✓ 已删除原文档中的图片")

        # 4. 输出摘要
        print("\n" + "=" * 80)
        print("✅ 处理完成!")
        print(f"  📝 DOCX: {output_docx_path.name}")
        print(f"  📄 PDF:  {output_pdf_path.name}")
        print(f"\n💡 使用提示:")
        print(f"  • PDF开头是图片索引目录（A3纵向，支持多页）")
        print(f"  • 有效图片: 连续编号 图1~图{active_count}")
        if orphan_count > 0:
            print(f"  • 孤儿图片: PDF末尾 (标注[未使用])")

        return True

    except Exception as e:
        print(f"❌ 处理失败: {e}")
        import traceback
        traceback.print_exc()
        return False


def get_docx_files_from_folder(folder_path):
    """获取文件夹下的所有 .docx 文件（不递归子文件夹）"""
    folder = Path(folder_path)
    if not folder.exists():
        print(f"❌ 路径不存在: {folder}")
        return []
    if not folder.is_dir():
        print(f"❌ 不是文件夹: {folder}")
        return []

    docx_files = []
    for item in sorted(folder.iterdir(), key=lambda p: p.name.lower()):
        if item.is_dir():
            continue
        if item.suffix.lower() != '.docx':
            continue
        # 跳过 Word 临时文件
        if item.name.startswith('~$'):
            continue
        # 跳过已做“图片标记”的输出文件，避免重复处理
        if '_已标记图片' in item.stem:
            print(f"    ⏭️  跳过已标记图片文件: {item.name}")
            continue
        docx_files.append(item)

    return docx_files


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description='DOCX 图片分离工具 - 增强版',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
功能特性:
  • 智能连续编号 - 只对实际使用的图片编号
  • PDF多页目录 - 在PDF开头生成图片索引（A3纵向）
  • 孤儿图片检测 - 识别未被引用的图片
  • 位置追踪 - 记录图片在文档中的位置

示例:
  python DOCX图片分离.py document.docx
  python DOCX图片分离.py document.docx --remove-images
  python DOCX图片分离.py document.docx --output-dir ./output/
  python DOCX图片分离.py ./docx_folder/ --output-dir ./output/
        """
    )

    parser.add_argument('input_path', help='DOCX 文件路径或包含 DOCX 的文件夹路径（不处理子文件夹）')
    parser.add_argument('--remove-images', '-r', action='store_true',
                        help='删除原文档中的图片（仅保留【图XXX】标记）')
    parser.add_argument('--output-dir', '-o', help='输出目录（默认为输入文件所在目录）')
    parser.add_argument('--no-optimize', action='store_true',
                        help='不优化图片格式（全部转PNG，文件会更大）')
    parser.add_argument('--jpeg-quality', type=int, default=85, metavar='Q',
                        help='JPEG质量 (1-100，默认85)')

    args = parser.parse_args()

    print("🚀 DOCX 图片分离工具 - 增强版")
    print("=" * 80)

    input_path = Path(args.input_path)

    if not input_path.exists():
        print(f"❌ 路径不存在: {input_path}")
        sys.exit(1)

    if input_path.is_dir():
        docx_files = get_docx_files_from_folder(input_path)
        if not docx_files:
            print(f"❌ 文件夹中未找到可处理的 .docx 文件: {input_path}")
            sys.exit(1)

        total = len(docx_files)
        success_count = 0
        fail_count = 0

        print(f"📂 批量处理文件夹: {input_path}")
        print(f"📄 待处理 DOCX: {total} 个 (不处理子文件夹)\n")

        for idx, docx_file in enumerate(docx_files, 1):
            print(f"\n{'=' * 80}")
            print(f"📄 [{idx}/{total}] {docx_file.name}")
            print(f"{'=' * 80}")

            ok = process_docx_file(
                docx_file,
                remove_images=args.remove_images,
                output_dir=args.output_dir,
                optimize_images=not args.no_optimize,
                jpeg_quality=args.jpeg_quality
            )
            if ok:
                success_count += 1
            else:
                fail_count += 1

        print(f"\n{'=' * 80}")
        print("📊 批量处理完成")
        print(f"  ✅ 成功: {success_count} 个")
        print(f"  ❌ 失败: {fail_count} 个")
        print(f"  📁 总计: {total} 个")

        success = fail_count == 0
    else:
        success = process_docx_file(
            input_path,
            remove_images=args.remove_images,
            output_dir=args.output_dir,
            optimize_images=not args.no_optimize,
            jpeg_quality=args.jpeg_quality
        )

    if success:
        print("\n✨ 全部完成!")
        sys.exit(0)
    else:
        print("\n💥 处理失败")
        sys.exit(1)


if __name__ == "__main__":
    main()
