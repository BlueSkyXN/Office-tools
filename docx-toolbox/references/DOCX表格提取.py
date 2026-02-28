#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX 表格提取工具

功能说明：
1. 识别 DOCX 文档中的所有表格
2. 在原文档中每个表格前插入【表N】标记，生成新的已标记文档
3. 提取所有表格内容为结构化数据
4. 导出为三种格式：
   - TXT 文件：使用标准 CSV 格式，每个表格用分隔线隔开
   - XLSX 文件：所有表格合并到"所有表格" sheet 中，每个表格有加粗标题
   - PDF 文件：智能自适应页面大小，包含可点击目录，每个表格独立页面

输出文件：
  - {原文件名}_已标记.docx    # 插入了【表N】标记的文档
  - {原文件名}_表格提取.txt   # CSV 格式的表格数据
  - {原文件名}_表格提取.xlsx  # Excel 格式的表格数据（单 sheet 包含所有表格）
  - {原文件名}_表格提取.pdf   # PDF 格式，自适应页面大小，含目录

PDF 特性：
  - 自动注册中文字体（支持 macOS 系统字体）
  - 智能计算每个表格的最佳页面尺寸（宽度和高度）
  - 首页为可点击的表格目录，显示：
    * 总页数（真实页数）
    * 每个表格的精确页码范围（如：第 2-5 页）
    * 每个表格的行数和占用页数（如：150 行，共 4 页）
  - 真实页码：使用两遍构建技术，目录显示的是渲染后的真实页码，非估算值
  - 自动分页：大表格（行数过多）自动跨页显示，避免内容丢失
  - 自动计算列宽，支持长文本换行
  - 表格左对齐，标题行灰色背景，首行在每页重复显示

用法：
  python DOCX表格提取.py <docx文件路径|文件夹路径>

示例：
  python DOCX表格提取.py document.docx
  python DOCX表格提取.py /path/to/docx_folder/

依赖：
  - python-docx: DOCX 文档读写
  - openpyxl: Excel 文件生成
  - reportlab: PDF 文件生成（含中文支持）

技术实现：
  - PageMarker 自定义 Flowable：用于在渲染过程中记录真实页码
  - 两遍构建技术：第一遍收集页码，第二遍生成最终 PDF
  - 目录一致性保证：第一遍使用与第二遍相同行数和长度的占位符，避免页码偏移
  - BaseDocTemplate + PageTemplate：实现每个表格独立页面尺寸
  - LongTable + repeatRows：自动分页，标题行重复显示
"""

import sys
import os
import csv
import io
from pathlib import Path
import argparse
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import openpyxl

# ReportLab imports for PDF generation
from reportlab.lib.pagesizes import A4, A3, landscape
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import BaseDocTemplate, PageTemplate, Frame, Table, LongTable, TableStyle, Paragraph, Spacer, PageBreak, NextPageTemplate, Flowable
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib.enums import TA_LEFT, TA_CENTER

TABLE_MARK_SUFFIX = "_已标记表格"

def insert_paragraph_before_table(table, text):
    """
    在表格前插入段落标记

    Args:
        table: python-docx Table 对象
        text: 要插入的文本内容（如 "【表1】"）

    实现原理：
        通过操作 DOCX 的底层 XML 结构，在表格元素前插入新的段落元素
    """
    # 获取表格的底层 XML 元素
    tbl_element = table._element
    parent = tbl_element.getparent()

    # 创建新段落的 XML 元素
    p = parse_xml(f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:t>{text}</w:t></w:r></w:p>')

    # 在表格元素之前插入段落
    parent.insert(parent.index(tbl_element), p)

def extract_table_data(table):
    """
    提取表格数据为二维列表

    Args:
        table: python-docx Table 对象

    Returns:
        list: 二维列表，每行是一个列表，包含该行所有单元格的文本

    数据清理：
        - 去除单元格首尾空白
        - 将单元格内的换行符替换为空格
    """
    data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            # 清理单元格文本：去除首尾空白，将换行符替换为空格
            cell_text = cell.text.strip().replace('\n', ' ')
            row_data.append(cell_text)
        data.append(row_data)
    return data

def save_to_txt(tables_data, output_path):
    """
    保存所有表格为 TXT 文件（CSV 格式）

    Args:
        tables_data: 列表，元素为 (表格编号, 表格数据) 的元组
        output_path: Path 对象，输出文件路径

    文件格式：
        【表1】
        "单元格1","单元格2",...
        "单元格1","单元格2",...
        ==================================================

        【表2】
        ...
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        for idx, data in tables_data:
            # 写入表格标题
            f.write(f"【表{idx}】\n")

            # 使用 csv 模块生成标准 CSV 格式字符串
            output = io.StringIO()
            writer = csv.writer(output)
            writer.writerows(data)
            f.write(output.getvalue())

            # 添加分隔线
            f.write("\n" + "="*50 + "\n\n")
    print(f"  ✓ 已导出 TXT: {output_path.name}")

def save_to_xlsx(tables_data, output_path):
    """
    保存所有表格为 XLSX 文件

    Args:
        tables_data: 列表，元素为 (表格编号, 表格数据) 的元组
        output_path: Path 对象，输出文件路径

    实现方式：
        - 创建一个名为"所有表格"的 sheet
        - 将所有表格依次写入，每个表格前有【表N】标题（加粗，12号字体）
        - 表格之间空两行

    备注：
        代码中包含注释的部分可选实现为每个表格创建单独的 sheet
    """
    wb = openpyxl.Workbook()
    # 删除 openpyxl 默认创建的 Sheet
    if 'Sheet' in wb.sheetnames:
        del wb['Sheet']

    # 创建"所有表格" Sheet，将所有表格合并写入
    ws_all = wb.create_sheet("所有表格")
    current_row = 1

    for idx, data in tables_data:
        # 写入表格标题（加粗显示）
        cell = ws_all.cell(row=current_row, column=1, value=f"【表{idx}】")
        cell.font = openpyxl.styles.Font(bold=True, size=12)
        current_row += 1

        # 逐行写入表格数据
        for row_data in data:
            for col_idx, cell_value in enumerate(row_data, 1):
                ws_all.cell(row=current_row, column=col_idx, value=cell_value)
            current_row += 1

        current_row += 2  # 表格间空两行

        # 可选实现：为每个表格创建独立的 Sheet
        # sheet_name = f"表{idx}"
        # ws = wb.create_sheet(sheet_name)
        # for r_idx, row_data in enumerate(data, 1):
        #     for c_idx, cell_value in enumerate(row_data, 1):
        #         ws.cell(row=r_idx, column=c_idx, value=cell_value)

    wb.save(output_path)
    print(f"  ✓ 已导出 XLSX: {output_path.name}")

class PageMarker(Flowable):
    """
    用于记录特定位置的真实页码的 Flowable

    在 PDF 渲染过程中，这个类会被插入到 story 中，
    当 ReportLab 渲染到这个位置时，会调用 draw() 方法，
    此时可以获取到当前的真实页码。
    """
    def __init__(self, key, page_tracker):
        """
        Args:
            key: 标记的唯一标识符（如 'table_1_start'）
            page_tracker: 用于存储页码的字典
        """
        Flowable.__init__(self)
        self.key = key
        self.page_tracker = page_tracker
        self.width = 0
        self.height = 0

    def draw(self):
        """
        渲染时被调用，记录当前页码
        """
        # canv.getPageNumber() 返回当前页码（从1开始）
        self.page_tracker[self.key] = self.canv.getPageNumber()

def register_chinese_font():
    """
    注册中文字体供 ReportLab PDF 生成使用

    Returns:
        str: 成功注册的字体名称，失败则返回 "Helvetica"

    实现逻辑：
        1. 按优先级尝试 macOS 系统常见的中文字体路径
        2. 找到第一个可用的字体文件后立即注册并返回
        3. 如果所有路径都失败，打印警告并返回后备字体 Helvetica

    字体优先级：
        - STHeiti Medium（黑体）
        - PingFang（苹方）
        - Songti（宋体）
        - STHeiti Light（细黑体）

    注意：
        Helvetica 不支持中文显示，如果返回此字体，PDF 中的中文会显示为方框或空白
    """
    # 尝试常见的 macOS 中文字体路径
    font_paths = [
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/Library/Fonts/Songti.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc"
    ]

    font_name = "CustomChinese"
    registered = False

    for path in font_paths:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(font_name, path))
                registered = True
                break
            except Exception:
                continue

    if not registered:
        print("⚠️  未找到合适的中文字体，PDF 中文可能无法显示。")
        return "Helvetica" # Fallback
    return font_name

def calculate_smart_col_widths(data, font_name, max_width=None, font_size=10):
    """
    智能计算表格列宽，根据内容自适应

    Args:
        data: 二维列表，表格数据
        font_name: 字体名称（用于宽度计算，当前版本未实际使用）
        max_width: 可选，页面最大可用宽度（points），用于压缩列宽
        font_size: 字体大小，默认 10 points

    Returns:
        list: 每列的宽度列表（单位：points）

    计算策略：
        1. 遍历所有单元格，计算每列的"期望宽度"（最长内容的宽度）
           - 中文字符按字号全宽计算
           - 大写字母按字号 0.8 倍计算
           - 小写字母和数字按字号 0.65 倍计算
           - 额外添加 16 points padding
        2. 单列最大宽度限制为 500 points（约 17.6cm），超出则强制换行
        3. 如果提供了 max_width 且总宽度超出：
           - 按比例压缩各列宽度
           - 保证每列至少为 min_col_width（4个字符宽度）

    注意：
        - 宽度计算是估算值，实际渲染可能有细微差异
        - 如果表格列数过多且页面太窄，可能无法满足 max_width 限制
    """
    if not data:
        return []

    num_cols = len(data[0])
    desired_widths = [0] * num_cols
    
    # 限制单列最大宽度，强制长文本换行
    # 500 points 约为 17.6cm，足够宽了
    MAX_SINGLE_COL_WIDTH = 500 
    
    # 1. 计算期望宽度
    for row in data:
        for i, cell in enumerate(row):
            if i < num_cols:
                # 估算文本宽度
                width = 0
                for char in str(cell):
                    if '\u4e00' <= char <= '\u9fff':
                        width += font_size # 中文全角
                    elif char.isupper():
                        width += font_size * 0.8 # 大写字母较宽
                    else:
                        width += font_size * 0.65 # 小写字母和数字
                # 加上 padding (左右各4 + 额外余量)
                width += 16 
                # 记录最大宽度，但不超过单列上限
                desired_widths[i] = max(desired_widths[i], min(width, MAX_SINGLE_COL_WIDTH))

    total_desired = sum(desired_widths)
    
    # 如果没有限制，或者在限制内，直接返回
    if max_width is None or total_desired <= max_width:
        return desired_widths
        
    # 3. 压缩策略
    # 设定最小列宽，防止压得太扁
    min_col_width = font_size * 4 # 至少容纳4个字
    
    # 检查最小总宽度
    min_total_width = num_cols * min_col_width
    
    # 如果最小总宽度都超过了 max_width，说明页面实在太窄了
    # 这种情况下，我们优先保证最小宽度，允许超出 max_width (由外层逻辑决定是否扩大页面)
    if min_total_width > max_width:
        return [max(w, min_col_width) for w in desired_widths]
    
    # 如果可以压缩到 max_width
    # 计算需要压缩的比例
    # 注意：不能简单乘比例，因为有些列可能已经很小了
    # 这里采用简单比例压缩，但设有下限
    
    available_space = max_width
    current_total = total_desired
    
    final_widths = []
    for w in desired_widths:
        ratio = w / current_total
        allocated = max_width * ratio
        final_widths.append(max(allocated, min_col_width))
        
    return final_widths

def calculate_optimal_page_size(data, font_name):
    """
    计算单个表格的最佳页面大小，使表格能完整显示而不截断

    Args:
        data: 二维列表，表格数据
        font_name: 字体名称

    Returns:
        tuple: (page_width, page_height)，单位为 points

    计算逻辑：
        宽度计算：
        1. 调用 calculate_smart_col_widths() 获取无限制下的期望列宽
        2. 总宽度 = sum(列宽) * 1.05（增加 5% 安全余量）
        3. 最小宽度为 842 points（A3 landscape 宽度）
        4. 最大宽度限制为 14000 points

        高度计算：
        1. 估算每个单元格可能的最大行数（基于文本长度和列宽）
        2. 计算最高单元格的高度（行数 * 行高 14）
        3. 页面高度 = 最高单元格高度 + 4 inch（上下边距和标题空间）
        4. 最小高度为 595 points（A4 高度）
        5. 最大高度限制为 14000 points

    使用场景：
        在 save_to_pdf() 中为每个表格创建独立的 PageTemplate，
        确保宽表格和高表格都能在一个页面内完整显示
    """
    if not data:
        return A4
        
    # 1. 宽度计算
    # 获取无限制的期望宽度
    widths = calculate_smart_col_widths(data, font_name, max_width=None)
    # 增加 5% 的安全余量，防止计算误差导致换行或截断
    table_width = sum(widths) * 1.05
    
    # 2. 高度计算 (估算最长单元格的高度)
    max_cell_height = 0
    for row in data:
        for i, cell in enumerate(row):
            if i < len(widths):
                col_w = widths[i]
                text_len = len(str(cell))
                # 估算行数: (文本长度 * 字号) / (列宽 - padding)
                # 假设平均字符宽度为 font_size * 0.8 (中英文混合)
                est_lines = (text_len * 10 * 0.8) / (col_w - 8) 
                est_lines = max(1, int(est_lines) + 1)
                cell_h = est_lines * 14 # leading=14
                max_cell_height = max(max_cell_height, cell_h)
    
    # 页面高度至少要能容纳这一行 + 上下边距 + 标题空间
    # 默认高度 A3 Landscape (842)
    default_h = 842
    required_h = max_cell_height + 4*inch 
    
    # 限制最大尺寸
    MAX_PAGE_WIDTH_LIMIT = 14000 
    MAX_PAGE_HEIGHT_LIMIT = 14000
    
    page_width = max(table_width + 2*inch, 842) # 至少 A4 宽 (其实是 A3 Landscape 宽)
    page_height = max(required_h, 595) # 至少 A4 高
    
    # 限制
    page_width = min(page_width, MAX_PAGE_WIDTH_LIMIT)
    page_height = min(page_height, MAX_PAGE_HEIGHT_LIMIT)
    
    return (page_width, page_height)

def save_to_pdf(tables_data, output_path):
    """
    保存所有表格为 PDF 文件，使用 BaseDocTemplate 实现每页自适应大小

    Args:
        tables_data: 列表，元素为 (表格编号, 表格数据) 的元组
        output_path: Path 对象，输出文件路径

    核心特性：
        1. 多页面模板：使用 BaseDocTemplate 和 PageTemplate 实现每个表格独立页面尺寸
        2. 自适应页面：每个表格根据内容自动计算最佳页面宽度和高度
        3. 中文支持：自动注册 macOS 系统中文字体
        4. 真实页码：使用两遍构建技术获取精确页码，目录显示真实页码而非估算值
        5. 可点击目录：首页包含所有表格的超链接目录，显示：
           - 总页数（真实值）
           - 每个表格的精确页码范围（单页或跨页）
           - 每个表格的行数和占用页数
        6. 智能分页：使用 LongTable 自动处理跨页表格，标题行在每页重复
        7. 智能列宽：自动计算列宽，支持长文本自动换行

    实现流程（两遍构建）：
        第一阶段 - 准备页面模板：
        1. 创建目录页模板（A4 尺寸）
        2. 为每个表格计算最佳页面尺寸
        3. 为每个表格创建独立的 PageTemplate（不同 pagesize）

        第二阶段 - 第一遍构建（分析页码）：
        1. 创建临时 PDF，包含占位符目录和所有表格
        2. 关键设计：占位符目录与第二遍目录保持相同的行数和长度
           - 为每个表格生成一行占位符
           - 使用最长格式（"第 999-999 页（X 行，共 99 页）"）
           - 确保两遍的目录页数完全一致，避免页码偏移
        3. 在每个表格的开始和结束位置插入 PageMarker（自定义 Flowable）
        4. 渲染过程中，PageMarker.draw() 被调用，记录真实页码到 page_tracker
        5. 完成渲染，获得每个表格的起始页和结束页

        第三阶段 - 第二遍构建（生成最终 PDF）：
        1. 使用第一遍收集的真实页码生成完整目录
        2. 目录显示每个表格的精确页码范围和行数：
           - 单页表格：【表N】第 X 页（Y 行）
           - 跨页表格：【表N】第 X-Z 页（Y 行，共 M 页）
        3. 添加所有表格（与第一遍相同，但不插入 PageMarker）
        4. 渲染最终 PDF
        5. 清理临时文件

    两遍构建的优势：
        - 页码完全精确，无估算误差
        - AI 和读者能准确知道每个表格的完整范围
        - 避免因单元格换行导致的页数偏差
        - 通过保持目录一致性，避免目录长度变化导致的页码偏移问题

    样式说明：
        - 字体：CustomChinese（macOS 系统中文字体）或 Helvetica（后备）
        - 字号：正文 10pt，标题 14pt，目录 12pt
        - 表格：左对齐，标题行灰色背景，黑色网格线（0.5pt）
        - 边距：每页 0.5 inch

    异常处理：
        如果 PDF 生成失败，打印错误信息但不中断程序
    """
    font_name = register_chinese_font()
    
    # 定义样式
    styles = getSampleStyleSheet()
    style_cn = ParagraphStyle(
        name='ChineseStyle',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        leading=14,
        wordWrap='CJK',
        alignment=TA_LEFT,
    )
    
    style_title = ParagraphStyle(
        name='TitleStyle',
        parent=styles['Heading1'],
        fontName=font_name,
        fontSize=14,
        leading=18,
        spaceAfter=12,
        alignment=TA_LEFT # 标题左对齐
    )

    style_toc = ParagraphStyle(
        name='TOCStyle',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=12,
        leading=16,
        spaceAfter=6
    )

    # --- 1. 准备页面模板 ---
    page_templates = []
    
    # 目录页模板 (A4)
    toc_frame = Frame(0.5*inch, 0.5*inch, A4[0]-inch, A4[1]-inch, id='toc_frame')
    page_templates.append(PageTemplate(id='TOC', frames=[toc_frame], pagesize=A4))
    
    # 为每个表格创建一个模板
    table_sizes = {} # 存储每个表格的计算尺寸
    
    for idx, data in tables_data:
        if not data:
            continue
            
        # 计算该表格的最佳页面大小
        p_w, p_h = calculate_optimal_page_size(data, font_name)
        table_sizes[idx] = (p_w, p_h)
        
        # 创建 Frame (留出边距)
        margin = 0.5 * inch
        frame_w = p_w - 2 * margin
        frame_h = p_h - 2 * margin
        
        frame = Frame(margin, margin, frame_w, frame_h, id=f'frame_{idx}')
        template = PageTemplate(id=f'PT_{idx}', frames=[frame], pagesize=(p_w, p_h))
        page_templates.append(template)

    # 创建文档对象
    doc = BaseDocTemplate(str(output_path), pageTemplates=page_templates)

    # --- 2. 第一遍构建：渲染表格并记录真实页码 ---
    print("  🔍 第一遍：分析表格页码...")

    page_tracker = {}  # 用于存储真实页码
    story_first = []

    # 目录页（第一遍用占位符，但保持与第二遍相同的行数，确保页数一致）
    # 关键设计：如果两遍的目录页数不同，会导致所有表格的页码偏移！
    # 例如：第一遍目录1页，表1在第2页；第二遍目录2页，表1实际在第3页，但显示第2页 ✗
    # 解决方案：第一遍生成与第二遍相同行数和长度的目录占位符
    story_first.append(NextPageTemplate('TOC'))
    story_first.append(Paragraph("表格目录 / Table of Contents（正在分析...）", style_title))
    story_first.append(Spacer(1, 0.5*inch))

    # 为每个表格生成一行占位符，确保目录行数与第二遍相同
    for idx, data in tables_data:
        if not data:
            continue
        num_rows = len(data)
        # 占位符使用最长可能的格式（999-999页，99页），确保：
        # 1. 行数相同（每个表格一行）
        # 2. 长度足够（第二遍不会因为更长而额外换行，导致页数增加）
        story_first.append(Paragraph(
            f"• 【表{idx}】............ 第 999-999 页（{num_rows} 行，共 99 页）",
            style_toc
        ))

    story_first.append(PageBreak())

    # 添加表格并记录页码
    for idx, data in tables_data:
        if not data:
            continue

        # 切换到该表格对应的页面模板
        story_first.append(NextPageTemplate(f'PT_{idx}'))

        # 记录表格起始页码
        story_first.append(PageMarker(f'table_{idx}_start', page_tracker))

        story_first.append(Paragraph(f"<a name='Table_{idx}'/>【表{idx}】", style_title))

        # 获取该页面的可用宽度
        p_w, p_h = table_sizes[idx]
        available_width = p_w - inch

        # 计算列宽
        col_widths = calculate_smart_col_widths(data, font_name, max_width=available_width)

        table_data = []
        for row in data:
            row_data = []
            for cell in row:
                p = Paragraph(str(cell), style_cn)
                row_data.append(p)
            table_data.append(row_data)

        t = LongTable(table_data, colWidths=col_widths, repeatRows=1, hAlign='LEFT', splitInRow=1)

        t.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('ALIGN', (0, 0), (-1, 0), 'LEFT'),
        ]))

        story_first.append(t)
        story_first.append(Spacer(1, 0.5*inch))

        # 记录表格结束页码
        story_first.append(PageMarker(f'table_{idx}_end', page_tracker))
        story_first.append(PageBreak())

    # 第一遍构建（生成到临时位置以收集页码）
    temp_output = str(output_path).replace('.pdf', '_temp.pdf')
    doc_first = BaseDocTemplate(temp_output, pageTemplates=page_templates)

    page_analysis_ok = True
    try:
        doc_first.build(story_first)
        print(f"  ✓ 页码分析完成，发现 {len(page_tracker) // 2} 个表格")
    except Exception as e:
        page_analysis_ok = False
        print(f"  ⚠️  页码分析失败，将生成不含真实页码的目录: {e}")

    # --- 3. 第二遍构建：生成最终 PDF ---
    print("  📝 第二遍：生成最终 PDF...")

    story = []

    # 切换到目录模板
    story.append(NextPageTemplate('TOC'))

    # 整理表格的页码范围
    table_page_ranges = []
    total_pages = 1
    if page_analysis_ok and page_tracker:
        for idx, data in tables_data:
            if not data:
                continue

            start_key = f'table_{idx}_start'
            end_key = f'table_{idx}_end'

            if start_key in page_tracker and end_key in page_tracker:
                start_page = page_tracker[start_key]
                end_page = page_tracker[end_key]
                num_rows = len(data)
                table_page_ranges.append((idx, start_page, end_page, num_rows))

        # 计算总页数（从 page_tracker 中获取最大页码）
        total_pages = max(page_tracker.values())

        story.append(Paragraph(f"表格目录 / Table of Contents（共 {total_pages} 页）", style_title))
    else:
        story.append(Paragraph("表格目录 / Table of Contents（无真实页码）", style_title))
    story.append(Spacer(1, 0.5*inch))

    # 为每个表格添加目录项
    if page_analysis_ok and table_page_ranges:
        # 显示真实页码范围和行数
        for idx, start_page, end_page, num_rows in table_page_ranges:
            if start_page == end_page:
                page_info = f"第 {start_page} 页（{num_rows} 行）"
            else:
                page_info = f"第 {start_page}-{end_page} 页（{num_rows} 行，共 {end_page - start_page + 1} 页）"

            story.append(Paragraph(
                f"• <a href='#Table_{idx}'>【表{idx}】</a> ............ {page_info}",
                style_toc
            ))
    else:
        # 无真实页码时，仅显示行数
        for idx, data in tables_data:
            if not data:
                continue
            num_rows = len(data)
            story.append(Paragraph(
                f"• <a href='#Table_{idx}'>【表{idx}】</a> ............ （{num_rows} 行）",
                style_toc
            ))

    story.append(PageBreak())

    # --- 4. 添加表格（第二遍，与第一遍相同但不插入 PageMarker） ---
    for idx, data in tables_data:
        if not data:
            continue

        # 切换到该表格对应的页面模板
        story.append(NextPageTemplate(f'PT_{idx}'))

        story.append(Paragraph(f"<a name='Table_{idx}'/>【表{idx}】", style_title))

        # 获取该页面的可用宽度
        p_w, p_h = table_sizes[idx]
        available_width = p_w - inch

        # 计算列宽
        col_widths = calculate_smart_col_widths(data, font_name, max_width=available_width)

        table_data = []
        for row in data:
            row_data = []
            for cell in row:
                p = Paragraph(str(cell), style_cn)
                row_data.append(p)
            table_data.append(row_data)

        t = LongTable(table_data, colWidths=col_widths, repeatRows=1, hAlign='LEFT', splitInRow=1)

        t.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('ALIGN', (0, 0), (-1, 0), 'LEFT'),
        ]))

        story.append(t)
        story.append(Spacer(1, 0.5*inch))
        story.append(PageBreak())

    try:
        doc.build(story)
        print(f"  ✓ 已导出 PDF: {output_path.name}")
    except Exception as e:
        print(f"❌ PDF 生成失败: {e}")
    finally:
        # 清理临时文件（无论成功失败）
        if os.path.exists(temp_output):
            try:
                os.remove(temp_output)
            except Exception:
                pass


def process_docx(docx_path):
    """
    处理 DOCX 文件，提取表格并生成多种格式的输出

    Args:
        docx_path: str 或 Path 对象，输入的 DOCX 文件路径

    处理流程：
        1. 验证文件存在性并打开文档
        2. 识别文档中的所有表格（使用 doc.tables）
        3. 提取每个表格的数据为二维列表
        4. 在原文档中每个表格前插入【表N】标记（操作 XML 结构）
        5. 保存标记后的文档为 {原文件名}_已标记.docx
        6. 导出表格数据为三种格式：
           - TXT：标准 CSV 格式，表格间用分隔线隔开
           - XLSX：单个 sheet 包含所有表格，带加粗标题
           - PDF：自适应页面大小，含可点击目录

    输出文件：
        - {原文件名}_已标记.docx    # DOCX 格式，表格前插入了标记
        - {原文件名}_表格提取.txt   # TXT 格式，CSV 编码
        - {原文件名}_表格提取.xlsx  # Excel 格式
        - {原文件名}_表格提取.pdf   # PDF 格式，自适应页面大小

    异常处理：
        - 文件不存在：打印错误信息并返回
        - 文档打开失败：打印错误信息并返回
        - 文档中无表格：打印警告信息并返回

    注意事项：
        - 先收集所有表格对象到列表，避免遍历时修改文档结构导致的问题
        - 表格编号从 1 开始（符合人类阅读习惯）
    """
    docx_path = Path(docx_path)

    # 验证文件存在
    if not docx_path.exists():
        print(f"❌ 文件不存在: {docx_path}")
        return False

    # 避免把文件夹当文件打开
    if docx_path.is_dir():
        print(f"❌ 输入是文件夹，不是 DOCX 文件: {docx_path}")
        return False

    already_marked_table = docx_path.stem.endswith(TABLE_MARK_SUFFIX)
    if already_marked_table:
        print(f"🏷️  检测到已标记表格文件，将直接导出（不重复插入【表N】）: {docx_path.name}")

    print(f"📄 处理文件: {docx_path.name}")

    # 尝试打开文档
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"❌ 无法打开文档: {e}")
        return False

    tables_data = []

    # 先将所有表格对象收集到列表中
    # 注意：这样做是为了避免在遍历过程中修改文档结构可能导致的问题
    tables = list(doc.tables)

    # 检查是否存在表格
    if not tables:
        print("⚠️  文档中没有找到表格")
        return None

    print(f"  📊 发现 {len(tables)} 个表格")

    # 处理每个表格
    for idx, table in enumerate(tables, 1):
        # 1. 提取表格数据
        data = extract_table_data(table)
        tables_data.append((idx, data))

        # 2. 在文档中的表格前插入标记（已标记表格文件不重复插入，避免出现多个【表N】）
        if not already_marked_table:
            insert_paragraph_before_table(table, f"【表{idx}】")
        print(f"    处理 表{idx} ({len(data)}行)")

    # 3. 保存插入标记后的 DOCX 文档（如果输入已是标记表格文件，则不再额外生成）
    if not already_marked_table:
        output_docx_path = docx_path.parent / f"{docx_path.stem}{TABLE_MARK_SUFFIX}.docx"
        doc.save(output_docx_path)
        print(f"  ✓ 已保存标记文档: {output_docx_path.name}")

    # 4. 导出表格数据为 TXT 和 XLSX 格式
    output_txt_path = docx_path.parent / f"{docx_path.stem}_表格提取.txt"
    output_xlsx_path = docx_path.parent / f"{docx_path.stem}_表格提取.xlsx"
    output_pdf_path = docx_path.parent / f"{docx_path.stem}_表格提取.pdf"

    save_to_txt(tables_data, output_txt_path)
    save_to_xlsx(tables_data, output_xlsx_path)
    save_to_pdf(tables_data, output_pdf_path)

    print("\n✅ 处理完成!")
    return True

def get_docx_files_from_folder(folder_path: Path, *, include_marked: bool = False):
    """获取文件夹下的所有 .docx 文件（不递归子文件夹）

    过滤规则：
      - 只处理 .docx
      - 跳过 Word 临时文件（~$ 开头）
      - 默认跳过脚本自身输出的 *_已标记表格.docx，避免重复处理
    """
    folder_path = Path(folder_path)
    if not folder_path.exists():
        print(f"❌ 路径不存在: {folder_path}")
        return []
    if not folder_path.is_dir():
        print(f"❌ 不是文件夹: {folder_path}")
        return []

    docx_files = []
    for item in sorted(folder_path.iterdir(), key=lambda p: p.name.lower()):
        if item.is_dir():
            continue
        if item.suffix.lower() != ".docx":
            continue
        if item.name.startswith("~$"):
            continue
        # 跳过已做"表格标记"的输出文件，避免重复处理
        if TABLE_MARK_SUFFIX in item.stem:
            print(f"    ⏭️  跳过已标记表格文件: {item.name}")
            continue
        docx_files.append(item)

    return docx_files


def process_batch(folder_path: Path, *, include_marked: bool = False):
    """批量处理文件夹中的 docx 文件（不递归子文件夹）"""
    folder_path = Path(folder_path)
    print(f"📂 批量处理文件夹: {folder_path}")
    print("    🔍 扫描文件夹...")

    docx_files = get_docx_files_from_folder(folder_path, include_marked=include_marked)
    if not docx_files:
        print(f"❌ 文件夹中未找到可处理的 .docx 文件: {folder_path}")
        return False

    total = len(docx_files)
    print(f"    📄 待处理 DOCX: {total} 个 (不处理子文件夹)\n")

    ok_count = 0
    skip_count = 0
    fail_count = 0

    for idx, docx_file in enumerate(docx_files, 1):
        print(f"\n{'=' * 80}")
        print(f"📄 [{idx}/{total}] {docx_file.name}")
        print(f"{'=' * 80}")

        try:
            result = process_docx(docx_file)
        except Exception as e:
            print(f"❌ 处理异常: {e}")
            result = False

        if result is True:
            ok_count += 1
        elif result is None:
            skip_count += 1
        else:
            fail_count += 1

    print(f"\n{'=' * 80}")
    print("📊 批量处理完成")
    print(f"  ✅ 成功: {ok_count} 个")
    print(f"  ⏭️  跳过: {skip_count} 个")
    print(f"  ❌ 失败: {fail_count} 个")
    print(f"  📁 总计: {total} 个")

    return fail_count == 0


def main():
    parser = argparse.ArgumentParser(
        description="DOCX 表格提取工具",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python DOCX表格提取.py document.docx
  python DOCX表格提取.py /path/to/docx_folder/
  python DOCX表格提取.py /path/to/docx_folder/ --include-marked
        """.strip(),
    )
    parser.add_argument("input_path", help="DOCX 文件路径或包含 DOCX 的文件夹路径（不处理子文件夹）")
    parser.add_argument(
        "--include-marked",
        action="store_true",
        help="(已废弃) 该脚本现在默认会处理所有 .docx；该参数保留仅为兼容旧命令。",
    )
    args = parser.parse_args()

    input_path = Path(args.input_path)
    if not input_path.exists():
        print(f"❌ 路径不存在: {input_path}")
        sys.exit(1)

    if input_path.is_dir():
        success = process_batch(input_path, include_marked=args.include_marked)
    else:
        if input_path.suffix.lower() != ".docx":
            print(f"❌ 输入文件不是 .docx: {input_path}")
            sys.exit(1)
        success = process_docx(input_path) is True

    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
